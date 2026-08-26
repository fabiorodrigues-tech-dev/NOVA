#!/usr/bin/env python3
"""
Gerador de Documentos Word (.docx) — Ecossistema NOVA
Módulo de Carreira: Exportação de Currículos e Cover Letters em formato Microsoft Word.

Padrões de Design:
- Tipografia: Calibri / Arial corporativo
- Margens: 2.54 cm (1 polegada - Padrão Harvard / ABNT)
- ATS-Friendly: Tabela e formatação em coluna única sem caixas de texto flutuantes
- Hierarquia Visual: Cores sóbrias (#1A2530, #2C3E50, #555555)
"""

import os
import sys
import re
import argparse
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def strip_emojis(text: str) -> str:
    """Remove emojis e caracteres gráficos não-ASCII."""
    custom_symbols = [
        "■", "▪", "▫", "🔹", "🔸", "📍", "📧", "📱", "💼", "💻", "🚀", "🌌",
        "🎙️", "🎙", "🎓", "🎯", "🛠️", "🛠", "🔍", "⚡", "📅", "📝", "📊",
        "💡", "⚪", "🟢", "🟡", "❌", "🌟", "✨", "🔗", "⭐", "🏷️", "🏷", "✉️"
    ]
    for sym in custom_symbols:
        text = text.replace(sym, "")
    emoji_pattern = re.compile(r'[\U00010000-\U0010ffff]', flags=re.UNICODE)
    text = emoji_pattern.sub('', text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\|\s*\|', '|', text)
    return text.strip()

def add_bottom_border(paragraph, color_hex="1A2530", size="6"):
    """Adiciona uma linha horizontal sutil abaixo do parágrafo no Word."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{size}" w:space="4" w:color="{color_hex}"/></w:pBdr>')
    pPr.append(pBdr)

def add_formatted_text(paragraph, text: str, base_font_size=Pt(10), base_color=RGBColor(44, 62, 80), default_bold=False, default_italic=False):
    """
    Processa formatação inline básica de markdown (**negrito**, *itálico*, `código`) e insere runs.
    """
    text = strip_emojis(text)
    # Remove markdown links [texto](url) -> texto
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    # Tokeniza partes de negrito e itálico
    tokens = re.split(r'(\*\*[^\*]+\*\*|\*[^\*]+\*|`[^`]+`)', text)
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = 'Calibri'
        run.font.size = base_font_size
        run.font.color.rgb = base_color
        run.bold = default_bold
        run.italic = default_italic

        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(30, 30, 30)
        else:
            run.text = token

def gerar_cv_docx(markdown_path: str, output_docx_path: str):
    """
    Converte um currículo Markdown em um documento Word (.docx) no padrão Harvard Tech ATS.
    """
    if not os.path.exists(markdown_path):
        raise FileNotFoundError(f"Arquivo não encontrado: {markdown_path}")

    with open(markdown_path, "r", encoding="utf-8") as f:
        content = f.read()

    os.makedirs(os.path.dirname(os.path.abspath(output_docx_path)), exist_ok=True)

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    lines = content.splitlines()
    header_done = False
    i = 0

    while i < len(lines):
        line = lines[i]
        raw_line = line.strip()

        if not raw_line or raw_line == "---":
            i += 1
            continue

        # Cabeçalho Principal (Nome)
        if not header_done and raw_line.startswith("# "):
            name_text = raw_line[2:].strip()
            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_name.paragraph_format.space_after = Pt(2)
            p_name.paragraph_format.space_before = Pt(0)
            run = p_name.add_run(strip_emojis(name_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 37, 48)

            i += 1
            # Subtítulo
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and lines[i].strip().startswith("**"):
                sub_text = lines[i].strip().replace("**", "")
                p_sub = doc.add_paragraph()
                p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_sub.paragraph_format.space_after = Pt(2)
                run_sub = p_sub.add_run(strip_emojis(sub_text))
                run_sub.font.name = 'Calibri'
                run_sub.font.size = Pt(10)
                run_sub.font.bold = True
                run_sub.font.color.rgb = RGBColor(41, 128, 185)
                i += 1

            # Contato
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and ("@" in lines[i] or "|" in lines[i]):
                contact_text = lines[i].strip()
                p_contact = doc.add_paragraph()
                p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_contact.paragraph_format.space_after = Pt(6)
                add_formatted_text(p_contact, contact_text, base_font_size=Pt(8.5), base_color=RGBColor(100, 100, 100))
                add_bottom_border(p_contact, color_hex="1A2530", size="8")
                i += 1

            header_done = True
            continue

        # Título de Seção (## RESUMO, ## COMPETÊNCIAS, etc.)
        if raw_line.startswith("## "):
            sec_text = raw_line[3:].strip()
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(8)
            p_sec.paragraph_format.space_after = Pt(2)
            p_sec.paragraph_format.keep_with_next = True
            run = p_sec.add_run(strip_emojis(sec_text).upper())
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 37, 48)
            add_bottom_border(p_sec, color_hex="BDC3C7", size="4")
            i += 1
            continue

        # Subtítulo de Cargo / Empresa (### ...)
        if raw_line.startswith("### "):
            role_text = raw_line[4:].strip()
            p_role = doc.add_paragraph()
            p_role.paragraph_format.space_before = Pt(5)
            p_role.paragraph_format.space_after = Pt(1)
            p_role.paragraph_format.keep_with_next = True
            add_formatted_text(p_role, role_text, base_font_size=Pt(9.5), base_color=RGBColor(26, 37, 48), default_bold=True)
            i += 1
            continue

        # Metadados de período (*Abril de 2026 – Atual | Recife, PE*)
        if raw_line.startswith("*") and raw_line.endswith("*") and len(raw_line) < 100:
            meta_text = raw_line[1:-1].strip()
            p_meta = doc.add_paragraph()
            p_meta.paragraph_format.space_after = Pt(2)
            p_meta.paragraph_format.keep_with_next = True
            run = p_meta.add_run(strip_emojis(meta_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(120, 120, 120)
            i += 1
            continue

        # Bullet point (- ... ou • ... ou * ...)
        if raw_line.startswith("- ") or raw_line.startswith("• ") or (raw_line.startswith("* ") and not raw_line.endswith("*")):
            bullet_text = raw_line[2:].strip()
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_after = Pt(2)
            p_bullet.paragraph_format.left_indent = Inches(0.2)
            add_formatted_text(p_bullet, bullet_text, base_font_size=Pt(9), base_color=RGBColor(44, 62, 80))
            i += 1
            continue

        # Parágrafo comum
        p_body = doc.add_paragraph()
        p_body.paragraph_format.space_after = Pt(4)
        p_body.paragraph_format.line_spacing = 1.15
        add_formatted_text(p_body, raw_line, base_font_size=Pt(9), base_color=RGBColor(44, 62, 80))
        i += 1

    doc.save(output_docx_path)
    print(f"✅ Documento Word do Currículo gerado em: {output_docx_path}")

def gerar_cover_letter_docx(markdown_path: str, output_docx_path: str):
    """
    Converte uma Cover Letter Markdown em um documento Word (.docx) timbrado e formal.
    """
    if not os.path.exists(markdown_path):
        raise FileNotFoundError(f"Arquivo não encontrado: {markdown_path}")

    with open(markdown_path, "r", encoding="utf-8") as f:
        content = f.read()

    os.makedirs(os.path.dirname(os.path.abspath(output_docx_path)), exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    lines = content.splitlines()
    header_done = False
    i = 0

    while i < len(lines):
        line = lines[i]
        raw_line = line.strip()

        if not raw_line or raw_line == "---":
            i += 1
            continue

        # Título da Carta (# ✉️ Carta de Apresentação ...)
        if not header_done and raw_line.startswith("# "):
            title_text = raw_line[2:].strip()
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_after = Pt(4)
            run = p_title.add_run(strip_emojis(title_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 37, 48)

            # Extrai dados do candidato e contato
            i += 1
            while i < len(lines) and lines[i].strip().startswith("**"):
                c_line = lines[i].strip()
                p_c = doc.add_paragraph()
                p_c.paragraph_format.space_after = Pt(1)
                add_formatted_text(p_c, c_line, base_font_size=Pt(9), base_color=RGBColor(80, 80, 80))
                i += 1

            # Linha divisória timbrada
            p_div = doc.add_paragraph()
            p_div.paragraph_format.space_after = Pt(12)
            add_bottom_border(p_div, color_hex="1A2530", size="8")
            header_done = True
            continue

        # Destinatário ou Assunto (**À Equipe...** ou **Assunto:...**)
        if raw_line.startswith("**À") or raw_line.startswith("**Assunto"):
            p_rec = doc.add_paragraph()
            p_rec.paragraph_format.space_after = Pt(4)
            add_formatted_text(p_rec, raw_line, base_font_size=Pt(10), base_color=RGBColor(26, 37, 48), default_bold=True)
            i += 1
            continue

        # Itens numerados (1. NOVA, 2. Sofia, etc.)
        if re.match(r'^\d+\.\s+', raw_line):
            num_match = re.match(r'^(\d+\.\s+)(.+)', raw_line)
            num_prefix = num_match.group(1)
            item_text = num_match.group(2)

            p_item = doc.add_paragraph()
            p_item.paragraph_format.left_indent = Inches(0.25)
            p_item.paragraph_format.space_after = Pt(4)
            p_item.paragraph_format.line_spacing = 1.15
            run_num = p_item.add_run(num_prefix)
            run_num.font.name = 'Calibri'
            run_num.font.size = Pt(9.5)
            run_num.font.bold = True
            run_num.font.color.rgb = RGBColor(41, 128, 185)
            add_formatted_text(p_item, item_text, base_font_size=Pt(9.5), base_color=RGBColor(44, 62, 80))
            i += 1
            continue

        # Saudação / Fechamento
        if raw_line.startswith("Prezada") or raw_line.startswith("Atenciosamente") or raw_line.startswith("**Fábio Rodrigues**"):
            p_salut = doc.add_paragraph()
            p_salut.paragraph_format.space_before = Pt(8)
            p_salut.paragraph_format.space_after = Pt(4)
            add_formatted_text(p_salut, raw_line, base_font_size=Pt(10), base_color=RGBColor(26, 37, 48))
            i += 1
            continue

        # Parágrafo padrão da carta
        p_body = doc.add_paragraph()
        p_body.paragraph_format.space_after = Pt(6)
        p_body.paragraph_format.line_spacing = 1.15
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_formatted_text(p_body, raw_line, base_font_size=Pt(10), base_color=RGBColor(44, 62, 80))
        i += 1

    doc.save(output_docx_path)
    print(f"✅ Documento Word da Cover Letter gerado em: {output_docx_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Conversor de Markdown para Word DOCX (NOVA)")
    parser.add_argument("--type", choices=["cv", "cover_letter"], default="cv", help="Tipo de documento a gerar")
    parser.add_argument("--input", required=True, help="Caminho do arquivo markdown de entrada")
    parser.add_argument("--output", required=True, help="Caminho do arquivo DOCX de saída")

    args = parser.parse_args()
    if args.type == "cover_letter":
        gerar_cover_letter_docx(args.input, args.output)
    else:
        gerar_cv_docx(args.input, args.output)
